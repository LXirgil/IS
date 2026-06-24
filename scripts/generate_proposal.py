from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os

CONTENT_JSON = os.path.join(os.path.dirname(__file__), 'proposal_content.json')
OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'proposals')
OUT_FILE = os.path.join(OUT_DIR, 'AI_Bowling_Master_Proposal.docx')


def ensure_outdir():
    os.makedirs(OUT_DIR, exist_ok=True)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level==1 else 14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_paragraphs(doc, text):
    for para in text.split('\n\n'):
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(12)


def make_doc():
    with open(CONTENT_JSON, 'r', encoding='utf-8') as f:
        content = json.load(f)

    doc = Document()
    doc.core_properties.title = content.get('title', '企画書')

    # Title
    h = doc.add_heading(content.get('title', ''), level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Executive summary
    add_heading(doc, 'エグゼクティブサマリ', level=1)
    add_paragraphs(doc, content.get('executive_summary', ''))

    # Sections
    sections = ['problem', 'solution', 'features', 'technical_spec', 'implementation_plan', 'timeline', 'business_model', 'budget', 'metrics', 'team', 'ask', 'appendix']
    titles = {
        'problem': '課題',
        'solution': '解決策',
        'features': '主要機能',
        'technical_spec': '技術仕様',
        'implementation_plan': '実施計画',
        'timeline': 'スケジュール',
        'business_model': 'ビジネスモデル',
        'budget': '想定コスト（概算）',
        'metrics': '評価指標 (KPI)',
        'team': '必要体制',
        'ask': '要望・次のアクション',
        'appendix': '付録'
    }

    for key in sections:
        if content.get(key):
            add_heading(doc, titles[key], level=1)
            add_paragraphs(doc, content[key])

    ensure_outdir()
    doc.save(OUT_FILE)
    print('Saved:', OUT_FILE)


if __name__ == '__main__':
    make_doc()
